"""
Generate 可行性承诺书 / Feasibility Commitment Letter for FBFS project.
Output: project/FBFS/out/交付物/可行性承诺书.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/home/chu2026/Documents/github/automotive-apqp-agent/project/FBFS/out/交付物/可行性承诺书.docx"

# ── Colour constants ──────────────────────────────────────────────────────────
BLUE_RGB   = RGBColor(68, 114, 196)
BLUE_HEX   = "4472C4"
YELLOW_HEX = "FFF2CC"
WHITE_HEX  = "FFFFFF"
LIGHT_BLUE_HEX = "DCE6F1"   # alternating row tint (optional, not used here)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_shading(cell, fill_hex):
    """Set cell background colour via XML shading element."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell, border_hex="000000", border_sz=4):
    """Apply thin borders on all four sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border_el = OxmlElement(f"w:{side}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), str(border_sz))
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), border_hex)
        tcBorders.append(border_el)
    tcPr.append(tcBorders)


def style_header_cell(cell, text, font_size=10):
    """Blue background, white bold text."""
    set_cell_shading(cell, BLUE_HEX)
    set_cell_borders(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_data_cell(cell, text, yellow=False, font_size=10, bold=False, center=False):
    """Regular data cell, optionally yellow-highlighted."""
    if yellow:
        set_cell_shading(cell, YELLOW_HEX)
    set_cell_borders(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_heading(doc, number, title_zh, title_en):
    """Add a bold section heading with blue background."""
    para = doc.add_paragraph()
    para.clear()
    run = para.add_run(f"{number}、{title_zh}（{title_en}）")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Blue paragraph shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), BLUE_HEX)
    pPr.append(shd)

    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.2)
    return para


def add_table(doc, headers, rows, col_widths=None, yellow_cols=None, yellow_cells=None):
    """
    Add a bordered table.
    yellow_cols: set of column indices to shade yellow in data rows.
    yellow_cells: set of (row_idx, col_idx) (0-based data rows) to shade yellow.
    """
    yellow_cols  = yellow_cols  or set()
    yellow_cells = yellow_cells or set()

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # Header row
    for ci, hdr in enumerate(headers):
        style_header_cell(table.cell(0, ci), hdr)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            is_yellow = (ci in yellow_cols) or ((ri, ci) in yellow_cells)
            style_data_cell(cell, cell_text, yellow=is_yellow)

    doc.add_paragraph()  # spacing after table
    return table


# ── Main document ─────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins: 2 cm all sides
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Arial"

    # ── TITLE ────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("可行性承诺书 / Feasibility Commitment Letter")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Stellantis KP1 A&B — 燃油供给管路 (FC00SAA78530)")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()

    # ── 一、基本信息 ─────────────────────────────────────────────────────────
    add_section_heading(doc, "一", "基本信息", "Basic Information")

    basic_rows = [
        ("客户 / Customer",                   "Stellantis",                                        False),
        ("项目代号 / Program",                  "KP1 A&B",                                           False),
        ("零件号 / Part No.",                   "FC00SAA78530",                                      False),
        ("零件名称 / Part Name",                "Fuel Supply Line（供油管路，滤清器至发动机）",       False),
        ("RFQ 编号 / RFQ No.",                  "[待填 / TBD]",                                      True),
        ("评审日期 / Review Date",              "2026-03-10",                                        False),
        ("供应商名称 / Supplier",               "[公司名称 - 待填]",                                 True),
        ("供应商集成等级 / SI Level",           "Level 2",                                           False),
    ]

    table1 = doc.add_table(rows=len(basic_rows), cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = [6, 10]
    for i, w in enumerate(col_w):
        for cell in table1.columns[i].cells:
            cell.width = Cm(w)

    for ri, (label, value, yellow) in enumerate(basic_rows):
        lc = table1.cell(ri, 0)
        vc = table1.cell(ri, 1)
        style_data_cell(lc, label, bold=True)
        style_data_cell(vc, value, yellow=yellow)

    doc.add_paragraph()

    # ── 二、总体可行性结论 ───────────────────────────────────────────────────
    add_section_heading(doc, "二", "总体可行性结论", "Overall Feasibility Conclusion")

    conclusion_para = doc.add_paragraph()
    conclusion_run = conclusion_para.add_run("☐ Go   ☒ 条件 Go（Conditional Go）   ☐ No-Go")
    conclusion_run.bold = True
    conclusion_run.font.size = Pt(12)

    cond_heading = doc.add_paragraph()
    cond_heading.add_run("前提条件（Conditions）:").bold = True

    conditions = [
        "UG/Teamcenter (TcAE) 3D能力确认（Hard Prerequisite — Stellantis不接受其他3D格式）",
        "CDS（零部件开发规格书）从客户获取",
        "车辆开发时间计划（VP/PS/SOP日期）从客户获取",
        "金属管内镀镍方案确认（自有 or 外协）",
        "关键测试委外实验室预约确认",
    ]
    for i, cond in enumerate(conditions, 1):
        p = doc.add_paragraph(f"{i}. {cond}", style="List Number")
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # ── 三、特殊特性确认 ─────────────────────────────────────────────────────
    add_section_heading(doc, "三", "特殊特性确认", "CC/SC Commitment")

    cc_headers = ["特性类型", "特性名称", "要求值", "我方承诺 Cpk", "管控方案"]
    cc_rows = [
        ("CC", "燃油系统完整性（碰撞后不泄漏）", "FMVSS 301全通过",    "N/A（OEM测试）",  "接头二次锁止设计"),
        ("CC", "泄漏密封性（液体管）",             "≥150 PSI",           "[待填]",          "100%在线泄漏测试（15μm VLD）"),
        ("CC", "静电耗散（防爆燃）",               "SAE J1645",          "[待填]",          "导电层设计 + 接地路径验证"),
        ("SC", "接头拔脱力（液体管，RT）",         "≥450N",              "[待填]",          "接头设计和管端成型管控"),
        ("SC", "脉冲耐久（300K/600K次）",          "R95C90, 45件",       "[待填]",          "委外测试（TÜV/SGS）"),
        ("SC", "金属管内镀镍",                     "Must（CTS强制）",    "[待填]",          "电镀供应商管控"),
    ]
    # Yellow cells in col 3 (index 3) for rows where value is [待填]
    yellow_cc = {(ri, 3) for ri, row in enumerate(cc_rows) if "[待填]" in row[3]}
    add_table(doc, cc_headers, cc_rows,
              col_widths=[2, 5, 3.5, 3, 4],
              yellow_cells=yellow_cc)

    # ── 四、质量目标承诺 ─────────────────────────────────────────────────────
    add_section_heading(doc, "四", "质量目标承诺", "Quality Target Commitment")

    qt_headers = ["质量指标", "目标", "我方承诺"]
    qt_rows = [
        ("CCP @3月市场故障率",   "0.00 PPM",          "[待确认]"),
        ("CCP @12月市场故障率",  "0.035 PPM",         "[待确认]"),
        ("ICP 初始客户感知",      "0",                 "[待确认]"),
        ("关键尺寸 Cpk",          "≥1.67（推测）",    "[待填]"),
        ("整车设计寿命",          "15年/150,000英里", "确认"),
    ]
    yellow_qt = {(ri, 2) for ri, row in enumerate(qt_rows) if "[待" in row[2]}
    add_table(doc, qt_headers, qt_rows,
              col_widths=[6, 5, 6.5],
              yellow_cells=yellow_qt)

    # ── 五、制造可行性声明 ───────────────────────────────────────────────────
    add_section_heading(doc, "五", "制造可行性声明", "Manufacturing Feasibility")

    mfg_items = [
        "工艺路线确认：[待工程师确认 — 尼龙管成型 + 金属管弯曲 + 快速接头装配]",
        "DFM 结论：[待工程师评估]",
        "产能/节拍确认：[待工程师评估]",
        "金属管内镀镍：[自有 / 外协 — 待确认]",
        "阻尼器供应商：[待确认 Tier 2 配合意愿]",
    ]
    for item in mfg_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"• {item}")
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ── 六、技术例外项 ───────────────────────────────────────────────────────
    add_section_heading(doc, "六", "技术例外项", "Technical Exceptions")

    ex_headers = ["#", "例外项", "客户要求", "状态", "处置方向"]
    ex_rows = [
        ("1", "UG/TcAE 3D能力",          "Native .prt / TcAE",       "🔵 待确认", "内部确认/若无则为 Hard Stop"),
        ("2", "金属管内镀镍",             "CTS强制要求",              "🔵 待确认", "外协电镀 or 预镀镍管采购"),
        ("3", "CDS未提供",                "TDR要求",                  "❌ 已要求客户提供", "向客户PM索要"),
        ("4", "车辆开发计划未提供",       "TDR要求",                  "❌ 已要求客户提供", "向客户PM索要"),
        ("5", "脉冲/振动测试",            "R95C90, 45件",             "🔵 待确认", "委外第三方实验室"),
        ("6", "材料预批准（SD-M0008/03）","Stellantis预批准清单",     "🔵 待确认", "确认管材供应商状态"),
    ]
    add_table(doc, ex_headers, ex_rows,
              col_widths=[1, 4, 4, 3.5, 5])

    # ── 七、风险声明 ─────────────────────────────────────────────────────────
    add_section_heading(doc, "七", "风险声明", "Risk Statement")

    risk_intro = doc.add_paragraph()
    risk_intro.add_run("以下为主要高级别风险项（详见风险提示清单.md）：").bold = True

    risk_headers = ["风险级别", "风险描述", "管控措施"]
    risk_rows = [
        ("高", "UG/TcAE能力缺失 — 为报价 Hard Stop",                      "立即确认内部能力"),
        ("高", "CDS/3D/开发计划缺失导致成本精度低（±30%）",              "立即向客户索要"),
        ("高", "塑料管材/接头预批准状态未知",                              "TKO前确认材料和接头供应商批准状态"),
        ("高", "CCP 0.00 PPM@3月极严目标",                                "从设计阶段建立防错和100%在线检测"),
    ]
    add_table(doc, risk_headers, risk_rows,
              col_widths=[2.5, 8, 7])

    # ── 八、资源与进度承诺 ───────────────────────────────────────────────────
    add_section_heading(doc, "八", "资源与进度承诺", "Resource & Timing Commitment")

    res_headers = ["里程碑 / Milestone", "状态/说明", "备注"]
    res_rows = [
        ("CFT 团队确认",    "[待 SDT 名册填写]",          "见 SDT团队名册.xlsx"),
        ("SOP 日期",        "[待客户提供车辆开发计划]",    "❌ 缺失"),
        ("PPAP 提交日期",   "[待SOP日期确定后推算]",       "🔵 待定"),
        ("DV 测试开始",     "[待VP日期确认]",               "🔵 待定"),
        ("TKO",             "[待客户提供]",                 "❌ 缺失"),
    ]
    yellow_res = {(ri, 1) for ri, row in enumerate(res_rows) if "[待" in row[1]}
    add_table(doc, res_headers, res_rows,
              col_widths=[5, 7, 5.5],
              yellow_cells=yellow_res)

    # ── 九、签字栏 ───────────────────────────────────────────────────────────
    add_section_heading(doc, "九", "签字栏", "Sign-off")

    signoff_headers = ["工程负责人", "质量负责人", "销售负责人", "管理层代表"]
    # Four rows: title, name, date, signature
    signoff_sub_rows = [
        ["姓名：", "姓名：", "姓名：", "姓名："],
        ["日期：", "日期：", "日期：", "日期："],
        ["签字：", "签字：", "签字：", "签字："],
    ]

    sign_table = doc.add_table(rows=1 + len(signoff_sub_rows), cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i in range(4):
        for cell in sign_table.columns[i].cells:
            cell.width = Cm(4.3)

    # Header row
    for ci, hdr in enumerate(signoff_headers):
        style_header_cell(sign_table.cell(0, ci), hdr)

    # Sub rows — all yellow
    for ri, sub_row in enumerate(signoff_sub_rows):
        for ci, text in enumerate(sub_row):
            cell = sign_table.cell(ri + 1, ci)
            set_cell_shading(cell, YELLOW_HEX)
            set_cell_borders(cell)
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(text)
            run.font.size = Pt(11)
            # Make cell taller for signatures
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            trPr = cell._tc.getparent().get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), "600")
            trPr.append(trHeight)

    doc.add_paragraph()

    # ── FOOTER NOTE ──────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.left_indent = Cm(0.3)

    # Yellow shading on footer paragraph
    pPr = footer_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), YELLOW_HEX)
    pPr.append(shd)

    footer_text = (
        "\u26a0\ufe0f \u9ec4\u8272\u5b57\u6bb5\u9700\u4eba\u5de5\u586b\u5199\u3002"
        "\u672c\u627f\u8bfa\u4e66\u57fa\u4e8e\u73b0\u6709\u5ba2\u6237\u6587\u4ef6"
        "\uff08SSTS KP1 Rev.00, 2022-03-09\uff09\u81ea\u52a8\u751f\u6210\uff0c"
        "\u90e8\u5206\u6570\u636e\uff08CDS\u3001\u5f00\u53d1\u8ba1\u5212\u30013D\u6570\u636e\uff09"
        "\u7f3a\u5931\uff0c\u5df2\u6807\u6ce8\u4e3a\u5f85\u5b9a\u3002"
        "\u603b\u4f53\u7ed3\u8bba\u4e3a\u201c\u6761\u4ef6 Go\u201d\uff0c"
        "\u524d\u63d0\u6761\u4ef6\u987b\u9010\u9879\u786e\u8ba4\u540e\u65b9\u53ef\u6700\u7ec8\u7b7e\u5b57\u63d0\u4ea4\u3002"
    )
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(10)
    footer_run.italic = True

    # ── SAVE ─────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
